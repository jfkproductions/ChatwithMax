from flask import Flask, render_template, request, jsonify
from dotenv import load_dotenv
import os
import uuid
import pathlib
import requests
from bs4 import BeautifulSoup
from werkzeug.utils import secure_filename
from langchain_openai import ChatOpenAI
from langchain.memory import ConversationBufferMemory
from langchain.agents import create_openai_functions_agent, AgentExecutor
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder
from langchain_community.chat_message_histories.upstash_redis import UpstashRedisChatMessageHistory
from langchain_community.tools.tavily_search import TavilySearchResults
from langchain.tools.retriever import create_retriever_tool
from langchain_community.document_loaders import WebBaseLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings
from langchain_community.vectorstores.faiss import FAISS
import docx
from PyPDF2 import PdfReader
from langchain.schema import Document  # Correct import for Document schema

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Load environment variables
load_dotenv()

# Set USER_AGENT from environment or use a default
USER_AGENT = os.getenv("USER_AGENT", "MyCustomUserAgent/1.0")

# Function to load or create a UUID for the session
def get_or_create_session_id(file_path="session_id.txt"):
    file = pathlib.Path(file_path)
    if file.exists():
        with open(file, 'r') as f:
            session_id = f.read().strip()
    else:
        session_id = str(uuid.uuid4())
        with open(file, 'w') as f:
            f.write(session_id)
    return session_id

# Load URL and TOKEN from .env file
URL = os.getenv("URL")
TOKEN = os.getenv("TOKEN")

# Get or create session ID
session_id = get_or_create_session_id()

# Initialize Redis-based chat message history with the session ID
history = UpstashRedisChatMessageHistory(
    url=URL, token=TOKEN, ttl=500, session_id=session_id
)

# Set up the conversation memory
memory = ConversationBufferMemory(
    memory_key="chat_history",
    return_messages=True,
    chat_memory=history,
)

# Function to get all URLs from the website
def get_all_urls(base_url):
    visited = set()
    to_visit = [base_url]
    all_urls = []

    headers = {
        "User-Agent": USER_AGENT
    }

    while to_visit:
        url = to_visit.pop(0)
        if url in visited:
            continue

        visited.add(url)
        response = requests.get(url, headers=headers)
        soup = BeautifulSoup(response.content, "html.parser")

        # Add current URL to list
        all_urls.append(url)

        # Find all links on the page
        for link in soup.find_all('a', href=True):
            full_url = requests.compat.urljoin(base_url, link['href'])
            if base_url in full_url and full_url not in visited:
                to_visit.append(full_url)

    return all_urls

# Function to extract text from a Word document
def extract_text_from_word(file_path):
    doc = docx.Document(file_path)
    text = [p.text for p in doc.paragraphs]
    return "\n".join(text)

# Function to extract text from a PDF document
def extract_text_from_pdf(file_path):
    text = ""
    with open(file_path, 'rb') as file:
        reader = PdfReader(file)
        for page in reader.pages:
            text += page.extract_text()
    return text

# Helper function to create the agent
def create_agent(tools):
    model = ChatOpenAI(
        model='gpt-3.5-turbo-1106',
        temperature=0.7
    )

    prompt = ChatPromptTemplate.from_messages([
        ("system", "You are a friendly assistant called Max."),
        MessagesPlaceholder(variable_name="chat_history"),
        ("human", "{input}"),
        MessagesPlaceholder(variable_name="agent_scratchpad")
    ])

    agent = create_openai_functions_agent(
        llm=model,
        prompt=prompt,
        tools=tools
    )

    return AgentExecutor(
        agent=agent,
        tools=tools,
        memory=memory
    )

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/initialize', methods=['POST'])
def initialize():
    try:
        docs = []
        tools = []  # Start with an empty toolset
        tool_name = "dynamic_search"
        tool_hint = ""

        # Handle website URL input
        if request.form.get('input_type') == 'website':
            website_url = request.form.get('website_url')
            if website_url:
                # Get all URLs from the specified base URL
                all_urls = get_all_urls(website_url)
                print(f"Found {len(all_urls)} URLs to process.")

                # Load content from all URLs
                for url in all_urls:
                    loader = WebBaseLoader(url)
                    loaded_docs = loader.load()
                    docs.extend(loaded_docs)
                    print(f"Loaded {len(loaded_docs)} documents from {url}")

                tool_name = "web_search"
                tool_hint = f"Use this tool when searching for information on {website_url}."

                # Add TavilySearchResults only when a website is used
                tools.append(TavilySearchResults())

        # Handle file upload (Word or PDF)
        elif request.form.get('input_type') in ['word', 'pdf']:
            file = request.files['file']
            filename = secure_filename(file.filename)
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(file_path)

            if filename.endswith('.docx'):
                text = extract_text_from_word(file_path)
                tool_name = "word_search"
                tool_hint = "Use this tool when searching for information in the uploaded Word document."
            elif filename.endswith('.pdf'):
                text = extract_text_from_pdf(file_path)
                tool_name = "pdf_search"
                tool_hint = "Use this tool when searching for information in the uploaded PDF document."
            else:
                return jsonify({"status": "failed", "message": "Unsupported file format."})

            # Wrap the extracted text in a Document object (LangChain schema)
            doc = Document(page_content=text, metadata={"source": filename})
            docs.append(doc)

        else:
            return jsonify({"status": "failed", "message": "No valid input provided."})

        if not docs:
            return jsonify({"status": "failed", "message": "No documents were loaded. Please check the input and try again."})

        # Split documents
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=200,
            chunk_overlap=20
        )
        splitDocs = splitter.split_documents(docs)
        print(f"Split documents into {len(splitDocs)} chunks.")

        # Create embeddings and vector store
        embedding = OpenAIEmbeddings()
        vectorStore = FAISS.from_documents(splitDocs, embedding=embedding)
        retriever = vectorStore.as_retriever(search_kwargs={"k": 5})  # Increase k for broader retrieval
        print("Created retriever from vector store.")

        # Create a dynamic retriever tool depending on the input type
        retriever_tools = create_retriever_tool(
            retriever,
            tool_name,
            tool_hint
        )
        tools.append(retriever_tools)  # Add the retriever tool to the tools list

        # Create the agent executor
        global agentExecutor
        agentExecutor = create_agent(tools)

        # Respond with tool hint
        return jsonify({"status": "initialized", "tool_hint": tool_hint})

    except Exception as e:
        print(f"Initialization error: {str(e)}")
        return jsonify({"status": "failed", "message": str(e)})

@app.route('/chat', methods=['POST'])
def chat():
    try:
        data = request.json
        user_input = data.get('user_input')

        # Check if the user is asking a general question or if it's related to the data loaded
        if agentExecutor:
            response = agentExecutor.invoke({
                "input": user_input
            })
        else:
            # Fallback to just using ChatGPT for general queries if no agentExecutor is set
            model = ChatOpenAI(model='gpt-3.5-turbo-1106', temperature=0.7)
            response = model.invoke({"input": user_input})

        return jsonify({"response": response["output"]})
    
    except Exception as e:
        print(f"Chat error: {str(e)}")
        return jsonify({"status": "failed", "message": str(e)})

if __name__ == '__main__':
    app.run(debug=True)
